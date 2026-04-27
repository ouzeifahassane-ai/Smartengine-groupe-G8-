from docx import Document
from docx.shared import Pt
import os

def create_dossier_conception():
    doc = Document()
    
    # Titre principal
    title = doc.add_heading('Dossier de Conception - Projet SmartEngine', 0)
    
    # Sprint 1
    doc.add_heading('Sprint 1 : Cadrage du Projet', level=1)
    
    doc.add_heading('1. Contexte métier de RavenStack', level=2)
    doc.add_paragraph(
        "RavenStack est un fournisseur de services SaaS (Software as a Service) en pleine croissance. "
        "Dans un marché B2B hautement concurrentiel, la rétention des clients est devenue la priorité stratégique numéro un. "
        "RavenStack fait face à un phénomène de 'churn' (résiliation) qui impacte son revenu récurrent mensuel (MRR). "
        "Pour pérenniser son activité, l'entreprise doit passer d'une posture réactive à une posture proactive."
    )
    
    doc.add_heading('2. Objectifs du projet smartEngine', level=2)
    doc.add_paragraph("Le projet smartEngine a pour but de concevoir et déployer une plateforme d'IA capable de :")
    doc.add_paragraph("• Identifier les signaux faibles de désengagement client.", style='List Bullet')
    doc.add_paragraph("• Prédire avec précision la probabilité de churn pour chaque compte.", style='List Bullet')
    doc.add_paragraph("• Fournir une interface visuelle (Dashboard) aux équipes de Customer Success.", style='List Bullet')
    doc.add_paragraph("• Automatiser des alertes via des workflows intelligents (n8n).", style='List Bullet')
    
    # Sprint 2
    doc.add_heading('Sprint 2 : Modélisation et Résultats', level=1)
    
    doc.add_heading('3. Modélisation prédictive', level=2)
    
    doc.add_heading('3.1 Algorithmes et Performance', level=3)
    p = doc.add_paragraph()
    p.add_run("Plusieurs algorithmes ont été évalués :").bold = True
    doc.add_paragraph("• Logistic Regression (Retenu) : AUC-ROC de 0.696. Interprétabilité forte.", style='List Bullet')
    doc.add_paragraph("• Random Forest : Testé, performances moindres sur le rappel.", style='List Bullet')
    doc.add_paragraph("• XGBoost : Écarté (libomp manquant dans l'environnement).", style='List Bullet')
    
    doc.add_heading('3.2 Stratégie d\'entraînement', level=3)
    doc.add_paragraph("• Split : 80% entraînement / 20% test avec stratification.")
    doc.add_paragraph("• Déséquilibre : Géré avec class_weight='balanced' (78% non-churn / 22% churn).")
    doc.add_paragraph("• Métriques : AUC-ROC, Precision, Recall, F1-Score.")
    
    doc.add_heading('3.3 Caractéristiques (Features) importantes', level=3)
    doc.add_paragraph("Les variables les plus influentes identifiées sont : days_since_last_login, usage_trend_3m, ratio_critical_tickets, avg_resolution_delay et seniority_months.")
    
    doc.add_heading('3.4 Seuils de risque et Actions', level=3)
    doc.add_paragraph("• Élevé (> 0.7) : Risque critique, intervention immédiate.")
    doc.add_paragraph("• Modéré (0.4 - 0.7) : Surveillance accrue.")
    doc.add_paragraph("• Faible (< 0.4) : Risque normal.")
    
    doc.add_heading('3.5 Limites et Biais potentiels', level=3)
    doc.add_paragraph("Des biais peuvent exister par industrie et par pays. Le modèle nécessite un réapprentissage régulier pour s'adapter aux évolutions du SaaS.")
    
    # RGPD & Outils
    doc.add_heading('4. Contraintes RGPD', level=1)
    doc.add_paragraph("Conformité à l'Article 22 (Décisions automatisées) : Le score est une aide à la décision, pas une sanction automatique. Minimisation des données et transparence algorithmique garanties.")
    
    doc.add_heading('5. Choix d\'outils justifiés', level=1)
    doc.add_paragraph("Python/pandas, scikit-learn, Streamlit, n8n, Gemini CLI.")
    
    # Footer
    doc.add_paragraph("\nDocument de conception - Mise à jour Sprint 2 - 27 Avril 2026")
    
    # Sauvegarde
    output_path = 'docs/dossier-conception.docx'
    doc.save(output_path)
    print(f"Fichier {output_path} généré avec succès.")

if __name__ == "__main__":
    create_dossier_conception()
